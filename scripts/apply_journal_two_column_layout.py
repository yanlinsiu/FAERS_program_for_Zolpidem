from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path(
    r"D:\program_FAERS\docs_or_reports\zolpidem_faers_manuscript_draft_cn_journal_20260601_submission_ready.docx"
)
OUT = Path(
    r"D:\program_FAERS\docs_or_reports\zolpidem_faers_manuscript_draft_cn_journal_20260604_two_column_layout.docx"
)


PAGE_WIDTH_CM = 21.0
LEFT_RIGHT_MARGIN_CM = 2.15
GUTTER_CM = 0.8
CONTENT_WIDTH_CM = PAGE_WIDTH_CM - 2 * LEFT_RIGHT_MARGIN_CM
COLUMN_WIDTH_CM = (CONTENT_WIDTH_CM - GUTTER_CM) / 2


def set_run_font(run, size=None, bold=None, east_asia="宋体", ascii_font="Times New Roman"):
    run.font.name = ascii_font
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def set_style_font(style, size, east_asia="宋体", ascii_font="Times New Roman", bold=None):
    style.font.name = ascii_font
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def set_columns(section, num=2, space_cm=GUTTER_CM):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(int(Cm(space_cm).twips)))


def set_table_width(table, width_cm):
    width_twips = int(Cm(width_cm).twips)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w_nodes = tbl_pr.xpath("./w:tblW")
    tbl_w = tbl_w_nodes[0] if tbl_w_nodes else None
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")

    grid_nodes = tbl.xpath("./w:tblGrid")
    tbl_grid = grid_nodes[0] if grid_nodes else None
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(tbl.index(tbl_pr) + 1, tbl_grid)
    elif tbl.index(tbl_grid) != tbl.index(tbl_pr) + 1:
        tbl.remove(tbl_grid)
        tbl.insert(tbl.index(tbl_pr) + 1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)

    col_count = max(len(table.rows[0].cells) if table.rows else 0, 1)
    col_width = int(width_twips / max(col_count, 1))
    for _ in range(col_count):
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(col_width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(col_width))
            tc_w.set(qn("w:type"), "dxa")
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6" if edge != "insideH" else "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    for edge in ("left", "right", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def paragraph_role(index, text):
    if index == 0:
        return "title"
    if index in {1, 2, 3, 4, 12, 13, 14}:
        return "meta"
    if text in {"摘要", "Abstract", "参考文献"} or re.match(r"^\d+\s+", text):
        return "heading1"
    if re.match(r"^\d+\.\d+\s+", text):
        return "heading2"
    if text.startswith(("图 ", "表 ")) or text.startswith(("Figure ", "Table ")):
        return "caption"
    if index >= 99:
        return "reference"
    return "body"


def main():
    doc = Document(SRC)

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(LEFT_RIGHT_MARGIN_CM)
        section.right_margin = Cm(LEFT_RIGHT_MARGIN_CM)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.2)
        section.header_distance = Cm(1.27)
        section.footer_distance = Cm(1.27)
        set_columns(section, 2, GUTTER_CM)

    set_style_font(doc.styles["Normal"], 8.5)
    normal_pf = doc.styles["Normal"].paragraph_format
    normal_pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal_pf.line_spacing = Pt(13)
    normal_pf.space_after = Pt(0)

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        role = paragraph_role(i, text)
        pf = p.paragraph_format

        if role == "title":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = None
            pf.line_spacing = 1.1
            pf.space_after = Pt(8)
            for run in p.runs:
                set_run_font(run, 14.5, True, east_asia="黑体")
        elif role == "meta":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = None
            pf.line_spacing = 1.05
            pf.space_after = Pt(2)
            for run in p.runs:
                set_run_font(run, 8.5, None, east_asia="仿宋")
        elif role == "heading1":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = None
            pf.keep_with_next = True
            pf.line_spacing = 1.15
            pf.space_before = Pt(5)
            pf.space_after = Pt(3)
            for run in p.runs:
                set_run_font(run, 10, True, east_asia="黑体")
        elif role == "heading2":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = None
            pf.keep_with_next = True
            pf.line_spacing = 1.1
            pf.space_before = Pt(3)
            pf.space_after = Pt(2)
            for run in p.runs:
                set_run_font(run, 9, True, east_asia="楷体")
        elif role == "caption":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = None
            pf.line_spacing = 1.1
            pf.space_before = Pt(3)
            pf.space_after = Pt(4)
            for run in p.runs:
                set_run_font(run, 7.5, None, east_asia="黑体")
        elif role == "reference":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.left_indent = Cm(0.45)
            pf.first_line_indent = Cm(-0.45)
            pf.line_spacing = 1.0
            pf.space_after = Pt(1)
            for run in p.runs:
                set_run_font(run, 7.5, None)
        else:
            if text:
                pf.first_line_indent = Cm(0.74)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(13)
            pf.space_after = Pt(0)
            for run in p.runs:
                if run.text:
                    set_run_font(run, 8.5, run.bold)

    for shape in doc.inline_shapes:
        if shape.width.cm > COLUMN_WIDTH_CM:
            ratio = COLUMN_WIDTH_CM / shape.width.cm
            shape.width = Cm(COLUMN_WIDTH_CM)
            shape.height = Cm(shape.height.cm * ratio)

    for idx, table in enumerate(doc.tables, start=1):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_width(table, COLUMN_WIDTH_CM)
        set_table_borders(table)
        font_size = 5.8 if len(table.columns) >= 6 else 7.0
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, font_size, True if row_idx == 0 else run.bold)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


SRC = Path(
    r"D:\program_FAERS\docs_or_reports\zolpidem_faers_manuscript_draft_cn_journal_20260604_two_column_layout.docx"
)
OUT = Path(
    r"D:\program_FAERS\docs_or_reports\zolpidem_faers_manuscript_draft_cn_journal_20260604_two_column_layout_title_author_abstract_fixed.docx"
)


def set_run_font(run, size, bold=False, east_asia="宋体", ascii_font="Times New Roman", superscript=None):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.bold = bold
    if superscript is not None:
        run.font.superscript = superscript
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)
    paragraph._p = paragraph._element = None


def set_para(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=None, before=0, after=0, line=1.0):
    paragraph.alignment = align
    pf = paragraph.paragraph_format
    pf.first_line_indent = first_indent
    pf.left_indent = None
    pf.right_indent = None
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line


def add_labelled_run(paragraph, label, text, label_size=8.5, text_size=8.5):
    label_run = paragraph.add_run(label)
    set_run_font(label_run, label_size, True, east_asia="黑体")
    body_run = paragraph.add_run(text)
    set_run_font(body_run, text_size, False, east_asia="宋体")


def main():
    doc = Document(SRC)
    paragraphs = doc.paragraphs

    title = paragraphs[0]
    clear_paragraph(title)
    set_para(title, before=0, after=8, line=1.05)
    title_text = "基于 FAERS 数据库的唑吡坦相关老年跌倒事件\n信号挖掘及报告特征分析"
    for idx, part in enumerate(title_text.split("\n")):
        if idx:
            title.add_run().add_break()
        run = title.add_run(part)
        set_run_font(run, 14.5, True, east_asia="黑体", ascii_font="Times New Roman")

    author = paragraphs[1]
    clear_paragraph(author)
    set_para(author, before=0, after=4, line=1.0)
    run = author.add_run("邵燕鳞")
    set_run_font(run, 9.0, False, east_asia="仿宋")
    sup = author.add_run("1")
    set_run_font(sup, 6.5, False, east_asia="仿宋", superscript=True)

    affiliation = paragraphs[2]
    clear_paragraph(affiliation)
    set_para(affiliation, before=0, after=6, line=1.0)
    run = affiliation.add_run("1. 中山大学（广东 深圳 518000）")
    set_run_font(run, 8.0, False, east_asia="仿宋")

    # Keep correspondence and contribution for now, but make them compact like journal footnote metadata.
    for idx in (3, 4):
        p = paragraphs[idx]
        set_para(p, before=0, after=2, line=1.0)
        for run in p.runs:
            set_run_font(run, 7.5, False, east_asia="仿宋")

    objective = paragraphs[6].text.replace("目的 ", "")
    methods = paragraphs[7].text.replace("方法 ", "")
    results = paragraphs[8].text.replace("结果 ", "")
    conclusion = paragraphs[9].text.replace("结论 ", "")
    keywords = paragraphs[10].text.replace("关键词 ", "")

    abstract = paragraphs[5]
    clear_paragraph(abstract)
    set_para(abstract, before=5, after=3, line=1.05)
    add_labelled_run(abstract, "【摘要】", "", label_size=8.5)
    add_labelled_run(abstract, "目的 ", objective)
    add_labelled_run(abstract, "方法 ", methods)
    add_labelled_run(abstract, "结果 ", results)
    add_labelled_run(abstract, "结论 ", conclusion)

    keyword_p = paragraphs[10]
    clear_paragraph(keyword_p)
    set_para(keyword_p, before=0, after=8, line=1.05)
    add_labelled_run(keyword_p, "【关键词】", keywords, label_size=8.5)

    # Remove the old separate objective/methods/results/conclusion paragraphs after reusing their text.
    for idx in sorted([6, 7, 8, 9], reverse=True):
        delete_paragraph(paragraphs[idx])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

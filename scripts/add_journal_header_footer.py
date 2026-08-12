from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path(
    r"D:\program_FAERS\docs_or_reports\zolpidem_faers_manuscript_draft_cn_journal_20260604_two_column_layout.docx"
)
OUT = Path(
    r"D:\program_FAERS\docs_or_reports\zolpidem_faers_manuscript_draft_cn_journal_20260604_two_column_layout_header_footer.docx"
)

CN_HEADER = "药物流行病学杂志 2024 年 9 月第 33 卷第 9 期"
EN_HEADER = "Chin J Pharmacoepidemiol, Sep. 2024, Vol. 33, No.9"
URL = "https://ywlxbx.whuznhmedj.com/"


def set_odd_even_headers(doc):
    settings = doc.settings._element
    existing = settings.xpath("./w:evenAndOddHeaders")
    if not existing:
        settings.append(OxmlElement("w:evenAndOddHeaders"))


def clear_block(block):
    for paragraph in list(block.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    for table in list(block.tables):
        table._element.getparent().remove(table._element)


def set_run_font(run, size=7.0, east_asia="宋体", ascii_font="Times New Roman", bold=False):
    run.font.name = ascii_font
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    text = OxmlElement("w:t")
    text.text = "961"
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)
    set_run_font(run, 7.0)


def add_aligned_table(block, left_text=None, right_text=None, left_page=False, right_page=False):
    table = block.add_table(rows=1, cols=2, width=Cm(16.7))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.xpath("./w:tblW")
    tbl_w = tbl_w[0] if tbl_w else OxmlElement("w:tblW")
    if tbl_w.getparent() is None:
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(int(Cm(16.7).twips)))
    tbl_w.set(qn("w:type"), "dxa")

    grid_nodes = tbl.xpath("./w:tblGrid")
    grid = grid_nodes[0] if grid_nodes else OxmlElement("w:tblGrid")
    if grid.getparent() is None:
        tbl.insert(tbl.index(tbl_pr) + 1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in (int(Cm(8.35).twips), int(Cm(8.35).twips)):
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)

    for idx, cell in enumerate(table.rows[0].cells):
        cell.width = Cm(8.35)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tcw = cell._tc.get_or_add_tcPr().tcW
        if tcw is None:
            tcw = OxmlElement("w:tcW")
            cell._tc.get_or_add_tcPr().append(tcw)
        tcw.set(qn("w:w"), str(int(Cm(8.35).twips)))
        tcw.set(qn("w:type"), "dxa")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT

    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    if left_text:
        run = left.add_run(left_text)
        set_run_font(run, 7.0, east_asia="黑体", ascii_font="Times New Roman")
    if right_text:
        run = right.add_run(right_text)
        set_run_font(run, 7.0, east_asia="黑体", ascii_font="Times New Roman")
    if left_page:
        run = left.add_run(" ")
        set_run_font(run, 7.0)
        add_page_field(left)
    if right_page:
        run = right.add_run(" ")
        set_run_font(run, 7.0)
        add_page_field(right)
    return table


def set_page_number_start(section, start=961):
    sect_pr = section._sectPr
    pg = sect_pr.xpath("./w:pgNumType")
    pg = pg[0] if pg else OxmlElement("w:pgNumType")
    if pg.getparent() is None:
        sect_pr.append(pg)
    pg.set(qn("w:start"), str(start))


def main():
    doc = Document(SRC)
    set_odd_even_headers(doc)

    for section in doc.sections:
        section.different_first_page_header_footer = False
        section.header_distance = Cm(1.27)
        section.footer_distance = Cm(1.27)
        set_page_number_start(section, 961)

        clear_block(section.header)
        clear_block(section.even_page_header)
        clear_block(section.footer)
        clear_block(section.even_page_footer)

        add_aligned_table(section.header, left_text=CN_HEADER, right_text=URL, right_page=True)
        add_aligned_table(section.even_page_header, left_text=URL, right_text=EN_HEADER, left_page=True)
        add_aligned_table(section.footer, left_text=URL)
        add_aligned_table(section.even_page_footer, right_text=URL)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs_or_reports" / "thesis_full_draft_v1.md"
ABSTRACTS = ROOT / "docs_or_reports" / "abstracts.md"
OUT = ROOT / "docs_or_reports" / "thesis_full_draft_v1.docx"


def set_run_font(run, size=None, bold=None, italic=None, color=None, east_asia="宋体"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_format(p, first_line=True):
    fmt = p.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def add_text_with_superscript(paragraph, text, size=12, bold=False):
    parts = re.split(r"(<sup>\[[^\]]+\]</sup>)", text)
    for part in parts:
        if not part:
            continue
        m = re.match(r"<sup>(\[[^\]]+\])</sup>", part)
        if m:
            run = paragraph.add_run(m.group(1))
            set_run_font(run, size=9, bold=False)
            run.font.superscript = True
        else:
            # Minimal markdown cleanup for emphasis markers left in references.
            part = part.replace("**", "").replace("*", "")
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=bold)


def clean_heading_text(text):
    text = text.strip()
    text = re.sub(r"（草稿[^）]*）", "", text)
    text = text.replace("（草稿）", "")
    return text.strip()


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def ensure_child(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_width(parent, tag, width_dxa):
    width = ensure_child(parent, tag)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(int(width_dxa)))


def apply_table_geometry(table, total_width_dxa):
    cols = len(table.columns)
    base = total_width_dxa // cols
    widths = [base] * cols
    widths[-1] += total_width_dxa - sum(widths)

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    set_width(tbl_pr, "w:tblW", total_width_dxa)
    indent = ensure_child(tbl_pr, "w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "0")
    layout = ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for col_idx, width in enumerate(widths):
        table.columns[col_idx].width = Twips(width)
    for row in table.rows:
        row.height = None
        for col_idx, cell in enumerate(row.cells):
            cell.width = Twips(widths[col_idx])
            set_width(cell._tc.get_or_add_tcPr(), "w:tcW", widths[col_idx])


def table_rows(block):
    rows = []
    for line in block:
        if re.match(r"^\|\s*-", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            text = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            add_text_with_superscript(p, text, size=9, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, "EDEDED")
    content_width = (
        doc.sections[-1].page_width.twips
        - doc.sections[-1].left_margin.twips
        - doc.sections[-1].right_margin.twips
    )
    apply_table_geometry(table, content_width)
    doc.add_paragraph()


def add_heading(doc, text, level, page_break=True):
    text = clean_heading_text(text)
    if page_break and level == 1 and (text.startswith("第") or text == "参考文献"):
        if len(doc.paragraphs) > 1:
            doc.add_page_break()
    p = doc.add_heading(level=min(level, 3))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    add_text_with_superscript(p, text, size={1: 16, 2: 14, 3: 12}.get(level, 12), bold=True)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 4)


def add_center_title(doc, text, size=16, space_after=18, east_asia="黑体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, east_asia=east_asia)
    return p


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word 中右键更新目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def setup_toc_styles(doc):
    styles = doc.styles
    for name, size, left in [
        ("TOC 1", 12, 0),
        ("TOC 2", 12, 0.75),
        ("TOC 3", 12, 1.5),
    ]:
        if name in styles:
            style = styles[name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(size)
            style.paragraph_format.left_indent = Cm(left)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_after = Pt(4)


def add_front_matter(doc):
    # Title page placeholder, then abstracts and a Word TOC field.
    add_center_title(doc, "基于 FAERS 数据库的老年人唑吡坦相关跌倒事件药物警戒研究", size=18, space_after=24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("硕士学位论文初稿")
    set_run_font(run, size=14, bold=True, east_asia="宋体")
    doc.add_page_break()

    lines = ABSTRACTS.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            add_center_title(doc, line[2:].strip(), size=16, space_after=18)
            i += 1
            continue
        p = doc.add_paragraph()
        is_keyword = line.startswith("关键词") or line.startswith("Key words")
        set_paragraph_format(p, first_line=not is_keyword)
        add_text_with_superscript(p, line, size=10.5, bold=False)
        i += 1
        if i < len(lines) and lines[i].startswith("# Abstract"):
            doc.add_page_break()
    doc.add_page_break()

    add_center_title(doc, "目  录", size=16, space_after=28)
    p = doc.add_paragraph()
    add_toc_field(p)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section)
    set_section_header_footer(body_section)
    restart_page_numbering(body_section, 1)


def configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)


def restart_page_numbering(section, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def set_section_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.text = "基于 FAERS 数据库的老年人唑吡坦相关跌倒事件药物警戒研究"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_run_font(run, size=9, color=(90, 90, 90))

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    footer.runs[-1]._r.append(fld_begin)
    footer.runs[-1]._r.append(instr)
    footer.runs[-1]._r.append(fld_end)
    footer.add_run(" 页")
    for run in footer.runs:
        set_run_font(run, size=9, color=(90, 90, 90))


def setup_document(doc):
    section = doc.sections[0]
    configure_section(section)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
    setup_toc_styles(doc)

    set_section_header_footer(section)


def build():
    doc = Document()
    setup_document(doc)
    add_front_matter(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()

    i = 0
    # The complete Markdown starts with a document title and an internal note.
    # Front matter is generated explicitly above, so skip those two lines.
    while i < len(lines) and (lines[i].startswith("# 基于 FAERS") or lines[i].startswith(">") or not lines[i].strip()):
        i += 1
    first_body_h1 = True
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith(">"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text_with_superscript(p, line.lstrip("> ").strip(), size=10)
            i += 1
            continue
        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, table_rows(block))
            continue
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            page_break = True
            if level == 1 and first_body_h1:
                page_break = False
                first_body_h1 = False
            add_heading(doc, m.group(2).strip(), level, page_break=page_break)
            i += 1
            continue
        p = doc.add_paragraph()
        # References should not use first-line indent.
        first_line = not re.match(r"^\[[0-9]+\]", line)
        set_paragraph_format(p, first_line=first_line)
        add_text_with_superscript(p, line, size=10.5 if first_line else 10)
        i += 1

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
